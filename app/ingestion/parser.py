from dataclasses import dataclass, field
from enum import Enum
from typing import Optional
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from pathlib import Path
from loguru import logger


class BlockType(Enum):
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    TABLE = "table"
    LIST_ITEM = "list_item"


@dataclass
class DocumentBlock:
    content: str
    block_type: BlockType
    page_number: int
    section_title: str = ""
    parent_section: str = ""
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    doc_id: str
    filename: str
    blocks: list[DocumentBlock] = field(default_factory=list)
    total_pages: int = 0
    title: str = ""
    author: str = ""

# Class to parse DOCX files
class DOCXParser:
    def __init__(self):
        pass

    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        # Placeholder for DOCX parsing logic
        logger.info(f"Parsing {file_path}")
        doc = ParsedDocument(doc_id=doc_id, filename=file_path)
        docx_doc = Document(file_path)
        doc.title = docx_doc.core_properties.title or ""
        doc.author = docx_doc.core_properties.author or ""

        current_section, parent_section = "", ""

        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            style = para.style.name.lower()

            if "heading" in style:
                current_section = text
                if int(style.replace("heading", "").strip()) <= 2:
                    parent_section = current_section

                doc.blocks.append(DocumentBlock(
                    content=text,
                    block_type=BlockType.HEADING,
                    page_number=0,  # DOCX doesn't have pages, can be set to 0 or estimated
                    section_title=current_section,
                    parent_section=parent_section,
                    metadata={"style": style}
                ))
            elif style.startswith("list"):
                doc.blocks.append(DocumentBlock(
                    content=text,
                    block_type=BlockType.LIST_ITEM,
                    page_number=0,
                    section_title=current_section,
                    parent_section=parent_section,
                    metadata={"style": style}
                ))
            else:
                doc.blocks.append(DocumentBlock(
                    content=text,
                    block_type=BlockType.PARAGRAPH,
                    page_number=0,
                    section_title=current_section,
                    parent_section=parent_section,
                    metadata={"style": style}
                ))

        for table in docx_doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):  # Skip empty rows
                    rows.append(cells)

            if len(rows) < 2:  # Skip empty/single-row tables
                continue

            table_text = "\n".join(["\t".join(cells) for cells in rows])

            doc.blocks.append(DocumentBlock(
                content=table_text,
                block_type=BlockType.TABLE,
                page_number=0,
                section_title=current_section,  # Don't leave these empty!
                parent_section=parent_section,
                metadata={"num_rows": len(rows) - 1, "num_cols": len(rows[0])}
            ))

    


class PDFParser:
    def __init__(self):
        self.heading_font_size_threshold = 14


    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        logger.info(f"Parsing {file_path}")
        doc = ParsedDocument(doc_id=doc_id, filename=file_path)
        pdf = fitz.open(file_path)
        doc.total_pages = len(pdf)
        doc.title = pdf.metadata.get("title", "") or ""
        doc.author = pdf.metadata.get("author", "") or ""


        current_section, parent_section = "", ""


        for page_num in range(len(pdf)):
            page = pdf[page_num]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] != 0: continue  # Skip non-text
                block_text, block_type, font_size = self._classify_block(block)
                if not block_text.strip(): continue


                if block_type == BlockType.HEADING:
                    if font_size > 18:
                        parent_section = block_text.strip()
                    current_section = block_text.strip()


                doc.blocks.append(DocumentBlock(
                    content=block_text.strip(),
                    block_type=block_type,
                    page_number=page_num + 1,
                    section_title=current_section,
                    parent_section=parent_section,
                    metadata={"font_size": font_size}
                ))
        pdf.close()
        self._extract_tables(file_path, doc)
        logger.info(f"Parsed {len(doc.blocks)} blocks from {doc.total_pages} pages")
        return doc
    
    def _classify_block(self, block) -> tuple[str, BlockType, float]:
        text_parts = []
        total_font_size = 0
        bold_count = 0
        total_spans = 0

        for line in block["lines"]:
            for span in line["spans"]:
                text_parts.append(span["text"])
                total_font_size += span["size"]
                if "bold" in span.get("font", "").lower():
                    bold_count += 1
                total_spans += 1

        block_text = " ".join(text_parts)
        avg_font_size = total_font_size / total_spans if total_spans > 0 else 0
        is_mostly_bold = bold_count > total_spans * 0.5 if total_spans else False

        if avg_font_size >= self.heading_font_size_threshold and is_mostly_bold:
            return block_text, BlockType.HEADING, avg_font_size
        elif block_text.strip().startswith(("•", "-", "–", "1.", "2.", "3.")):
            return block_text, BlockType.LIST_ITEM, avg_font_size
        return block_text, BlockType.PARAGRAPH, avg_font_size
    
    def _extract_tables(self, file_path: str, doc: ParsedDocument):
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for table in tables:
                    table_text = "\n".join(["\t".join(str(cell or "").strip() for cell in row)
                                            for row in table if any(row)])
                    doc.blocks.append(DocumentBlock(
                        content=table_text,
                        block_type=BlockType.TABLE,
                        page_number=page_num + 1,
                        section_title="",
                        parent_section=""
                    ))



class DocumentParser:
    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()


    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf": return self.pdf_parser.parse(file_path, doc_id)
        elif ext in (".docx", ".doc"): return self.docx_parser.parse(file_path, doc_id)
        else: raise ValueError(f"Unsupported: {ext}")

